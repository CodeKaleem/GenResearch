# ============================================================
# GenResearch — DOCX Export Service
# Generates the Draft and Completion Guide .docx files
# ============================================================
from __future__ import annotations

import re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _clean_markdown(text: str) -> str:
    """Very basic markdown cleanup for word export."""
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("### ", "")
    text = text.replace("## ", "")
    text = text.replace("# ", "")
    return text

def _format_citation(entry: dict, style: str) -> str:
    """Format a single citation registry entry according to the selected style."""
    authors = entry.get("authors", "Unknown")
    year = entry.get("year", "n.d.")
    title = entry.get("title", "Untitled")
    url = entry.get("url", "")
    
    if style == "apa":
        return f"{authors} ({year}). {title}. {url}"
    elif style == "ieee":
        return f"{authors}, \"{title},\" {year}. [Online]. Available: {url}"
    elif style == "mla":
        return f"{authors}. \"{title}.\" {year}. {url}"
    elif style == "chicago":
        return f"{authors}. \"{title}.\" {year}. {url}"
    
    return f"{authors} ({year}). {title}. {url}"

def generate_draft_docx(
    output_path: str,
    draft_text: str,
    citation_registry: list[dict],
    citation_style: str,
    title: str,
):
    """Generate the main research paper draft .docx file."""
    doc = docx.Document()

    # Title
    doc.add_heading(title, 0)

    # Process draft text and replace inline citations [CR-001]
    # We will just do a basic string replacement or add a bibliography section
    
    # First, let's collect which citations were actually used
    used_cids = set(re.findall(r'\[CR-\d{3}\]', draft_text))
    
    # Optional: We could replace [CR-XXX] with (Author, Year) for APA in-text, 
    # but for now, we'll keep the IDs or just append the bibliography.
    # Let's do a simple formatting:
    
    paragraphs = draft_text.split('\n\n')
    for p in paragraphs:
        if p.strip():
            # If it's a heading in markdown
            if p.startswith('### '):
                doc.add_heading(p.replace('### ', ''), level=3)
            elif p.startswith('## '):
                doc.add_heading(p.replace('## ', ''), level=2)
            elif p.startswith('# '):
                doc.add_heading(p.replace('# ', ''), level=1)
            else:
                clean_p = _clean_markdown(p)
                doc.add_paragraph(clean_p)

    # References Section
    doc.add_page_break()
    doc.add_heading("References", level=1)
    
    # Sort used citations (or all registry citations if preferred)
    registry_dict = {e['id']: e for e in citation_registry}
    
    for cid in sorted(list(used_cids)):
        if cid[1:-1] in registry_dict: # [CR-001] -> CR-001
            entry = registry_dict[cid[1:-1]]
            formatted = _format_citation(entry, citation_style)
            p = doc.add_paragraph()
            p.add_run(f"{cid} ").bold = True
            p.add_run(formatted)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(output_path)


def generate_completion_guide_docx(
    output_path: str,
    topic: str,
    sufficiency_report: dict,
    flagged_items: list[dict],
    citation_verification: dict,
    section_critique: dict,
    final_qa: dict,
):
    """Generate the Completion Guide .docx file outlining remaining user tasks."""
    doc = docx.Document()
    
    doc.add_heading(f"Completion Guide: {topic}", 0)
    
    doc.add_paragraph(
        "This guide contains the automated assessment of your generated research draft, "
        "highlighting areas that require your manual review and completion."
    )
    
    # 1. Flagged Items (from retries)
    doc.add_heading("1. Action Items & Flagged Issues", level=1)
    if not flagged_items:
        doc.add_paragraph("No critical issues flagged during generation.", style="Intense Quote")
    else:
        for item in flagged_items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{item.get('node', 'Unknown')}] ").bold = True
            p.add_run(item.get('issue', ''))
            if item.get('action_required'):
                p.add_run(f"\nAction Required: ").italic = True
                p.add_run(item.get('action_required'))

    # 2. Sufficiency Review
    doc.add_heading("2. Source Coverage Analysis", level=1)
    sections = sufficiency_report.get('sections', {})
    if sections:
        for sec_name, sec_data in sections.items():
            conf = sec_data.get('confidence', 'unknown') if isinstance(sec_data, dict) else sec_data
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{sec_name}: ").bold = True
            p.add_run(f"Confidence - {conf.upper()}")
            if isinstance(sec_data, dict) and 'reasoning' in sec_data:
                p.add_run(f" ({sec_data['reasoning']})")
    else:
        doc.add_paragraph("No section coverage data available.")

    # 3. Citation Verification
    doc.add_heading("3. Citation & Factual Grounding", level=1)
    score = citation_verification.get('coverage_score', 0)
    doc.add_paragraph(f"Score: {score:.0%} claims verified.")
    unverified = citation_verification.get('unverified_claims', [])
    if unverified:
        doc.add_paragraph("The following claims require manual citation:", style="Heading 3")
        for uc in unverified:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(uc.get('claim', '')).bold = True
            p.add_run(f" ({uc.get('location', '')})")
    
    # 4. Writing Quality Critique
    doc.add_heading("4. Section Quality Critique", level=1)
    doc.add_paragraph(f"Overall Score: {section_critique.get('overall_score', 0)}/10")
    for sec_name, sec_data in section_critique.get('sections', {}).items():
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{sec_name}: ").bold = True
        p.add_run(f"Score {sec_data.get('average_score', 0)}/10")
        if sec_data.get('issues'):
            p.add_run(f" - Issues: {', '.join(sec_data['issues'])}")
            
    # 5. Final QA
    doc.add_heading("5. Whole-Paper Assessment", level=1)
    doc.add_paragraph(f"Final Score: {final_qa.get('overall_score', 0)}/10")
    if final_qa.get('issues'):
        for issue in final_qa.get('issues', []):
            p = doc.add_paragraph(style="List Bullet")
            desc = issue if isinstance(issue, str) else issue.get('description', '')
            p.add_run(desc)

    doc.save(output_path)
